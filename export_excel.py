# -*- coding: utf-8 -*-
# @Author: wujiyu115
# @Date:   2016-07-12 18:56:20
# @Last Modified by:   wujiyu115
# @Last Modified time: 2016-07-12 19:50:54


from excel_util import ExcelNoTemplateUtils

from docx import Document
from docx.shared import Inches
from db_info import MySqlDbData

c = MySqlDbData()

def export_xls():
	entu = ExcelNoTemplateUtils('dbdict.xls')
	heading_xf = entu.ezxf('font: bold on; align: wrap off, vert centre, horiz center')
	title_xf = entu.ezxf('font: bold on, color blue; align: wrap on, vert bottom, horiz center; border: left thin, right thin, top thin, bottom thin')
	data_xfs = entu.ezxf('font: bold off; align: wrap off, vert centre, horiz left; border: left thin, right thin, top thin, bottom thin')

	headings = ['Field', 'Type', 'Null', 'Key', 'Default', 'Extra']
	sheets = c.getDB() # 每个数据库作为一个sheet表格
	datas = {}
	for sheet in sheets:
		datas =c.getDatas(sheet)
		sheet = entu.addSheet(sheet)
		entu.write(sheet, headings, datas, heading_xf, title_xf, data_xfs)

	entu.save()

def export_doc():

	document = Document()
	document.add_heading('database  schema',0)
	# p = document.add_paragraph('export database schema')

	sheet="bee_char"
	recordset =c.getDatas(sheet)
	for db_name,items in recordset.iteritems():
		document.add_paragraph()
		document.add_paragraph(
		   db_name, style='Caption'
		)
		headings = ['Field', 'Type', 'Null', 'Key', 'Default', 'Extra']
		head_len = len(headings)
		table = document.add_table(rows=1, cols=head_len)
		table.style = 'Light Grid'
		hdr_cells = table.rows[0].cells
		for x in xrange(0,head_len):
			hdr_cells[x].text = headings[x]
		for item in items:
			row_cells = table.add_row().cells
			for x in xrange(0,head_len):
				row_cells[x].text = str(item[x])

	document.save('database.docx')

if __name__ == '__main__':
	print "==export start=="
	# export_xls()
	export_doc()
	print "==End=="
