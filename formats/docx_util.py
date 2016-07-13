# -*- coding: utf-8 -*-
# @Author: wujiyu115
# @Date:   2016-07-13 14:23:29
# @Last Modified by:   wujiyu115
# @Last Modified time: 2016-07-13 14:55:35


from docx import Document
from docx.shared import Inches



class DocUtil():

	def __init__(self, filename):
		self.filename = filename
		self.document = Document()


	def write(self,  headings, data):
		head_len = len(headings)
		self.document.add_heading('database  schema',0)
		for db_name,items in data.iteritems():
			self.document.add_paragraph()
			self.document.add_paragraph(
			   db_name, style='Caption'
			)
			table = self.document.add_table(rows=1, cols=head_len)
			table.style = 'Light Grid'
			hdr_cells = table.rows[0].cells
			for x in xrange(0,head_len):
				hdr_cells[x].text = headings[x]
			for item in items:
				row_cells = table.add_row().cells
				for x in xrange(0,head_len):
					row_cells[x].text = str(item[x])

	def save(self):
		self.document.save(self.filename)