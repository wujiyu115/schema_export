# -*- coding: utf-8 -*-
# @Author: wujiyu115
# @Date:   2016-07-13 14:23:29
# @Last Modified by:   wujiyu115
# @Last Modified time: 2016-07-15 10:15:29


from docx import Document
from docx.shared import Inches
from base_format import BaseFormat


class DocFormat(BaseFormat):

	def __init__(self, filename,headings):
		self.suffix = "doc" #后缀
		super(DocFormat, self).__init__(filename,headings)
		self.document = Document()


	def write(self,   data):
		super(DocFormat, self).write(data)
		headings =self.headings
		head_len = len(headings)
		self.document.add_heading(u'数据库结构文档',0)
		for db_name,table_info in data.iteritems():
			table_caption = db_name
			# print(table_info.table_desc)
			if table_info.table_desc   :
				table_caption = '%s(%s)'%(db_name,table_info.table_desc)
			self.document.add_paragraph()
			self.document.add_paragraph(
			   table_caption, style='Caption'
			)
			table = self.document.add_table(rows=0, cols=head_len)
			table.style = 'Light Grid'
			# hdr_cells = table.rows[0].cells
			# for x in xrange(0,head_len):
			# 	hdr_cells[x].text = headings[x]
			for item in table_info.columns or []:
				row_cells = table.add_row().cells
				for x in xrange(0,head_len):
					txt =  item[x]
					if txt is None:
						txt = ""
					row_cells[x].text = txt

	def save(self):
		self.document.save(self.outname)